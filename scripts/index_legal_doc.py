"""
Index Luật-36-2024-QH15.docx into Qdrant collection "llm".
- No MariaDB writes
- No title field in payload — only content + source metadata
- Chunk size: 500 tokens, overlap: 50 tokens, separator: newline
"""
import os
import sys
import uuid
import logging

import docx
from llama_index.core import Document
from llama_index.core.node_parser import TokenTextSplitter
from openai import OpenAI
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

logging.basicConfig(level=logging.INFO, format="%(levelname)s %(message)s")
logger = logging.getLogger(__name__)

DOCX_PATH = os.path.join(os.path.dirname(__file__), "..", "data", "Luật-36-2024-QH15.docx")
COLLECTION = "llm"
QDRANT_URL = "http://localhost:6333"
EMBED_MODEL = "text-embedding-3-large"
EMBED_DIM = 3072
CHUNK_SIZE = 500   # tokens — preserves full article clauses for Vietnamese legal text
CHUNK_OVERLAP = 50

openai_client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
qdrant = QdrantClient(url=QDRANT_URL)


def read_docx(path: str) -> str:
    doc = docx.Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def chunk_text(text: str) -> list[str]:
    doc = Document(text=text)
    splitter = TokenTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separator="\n",
    )
    nodes = splitter.get_nodes_from_documents([doc])
    return [n.get_content() for n in nodes]


def embed(texts: list[str]) -> list[list[float]]:
    response = openai_client.embeddings.create(model=EMBED_MODEL, input=texts)
    return [item.embedding for item in response.data]


def ensure_collection():
    existing = {c.name for c in qdrant.get_collections().collections}
    if COLLECTION not in existing:
        logger.info("Creating collection %s", COLLECTION)
        qdrant.create_collection(
            collection_name=COLLECTION,
            vectors_config=VectorParams(size=EMBED_DIM, distance=Distance.DOT),
        )
    else:
        logger.info("Collection %s already exists", COLLECTION)


def index(chunks: list[str]):
    batch_size = 10
    point_id = 0
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        vectors = embed(batch)
        points = [
            PointStruct(
                id=point_id + j,
                vector=vectors[j],
                payload={"content": batch[j], "source": "Luật-36-2024-QH15", "page": ""},
            )
            for j in range(len(batch))
        ]
        qdrant.upsert(collection_name=COLLECTION, wait=True, points=points)
        point_id += len(batch)
        logger.info("Indexed chunks %d–%d / total so far: %d", i, i + len(batch) - 1, point_id)


def main():
    logger.info("Reading %s", DOCX_PATH)
    text = read_docx(DOCX_PATH)
    logger.info("Document length: %d chars", len(text))

    chunks = chunk_text(text)
    logger.info("Chunks produced: %d (size=%d tokens, overlap=%d)", len(chunks), CHUNK_SIZE, CHUNK_OVERLAP)

    ensure_collection()
    index(chunks)
    logger.info("Done. Total points indexed: %d", len(chunks))


if __name__ == "__main__":
    main()
